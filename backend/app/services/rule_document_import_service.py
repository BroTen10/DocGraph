"""规则文档导入服务：从上传的规则描述文档（PDF/EXCEL/WORD/MD）中提取文本，
然后复用 rule_import_service.import_rules_from_text 调用 LLM 解析为结构化规则。

支持格式：
- PDF: pdfplumber 提取文本（与 OCR 服务共用）
- Excel (.xlsx/.xls): openpyxl 读取所有 sheet 的单元格文本
- Word (.docx): python-docx 提取段落 + 表格文本
- Markdown (.md): 直接读取文本
"""

from __future__ import annotations

import logging
import os
import tempfile
import uuid
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session

from .rule_import_service import import_rules_from_text

logger = logging.getLogger(__name__)


# 允许的文件扩展名
ALLOWED_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".docx", ".md", ".txt"}


def _get_file_ext(filename: str) -> str:
    """获取文件扩展名（小写）。"""
    return Path(filename).suffix.lower()


def _validate_file(filename: str) -> str:
    """校验文件类型，返回扩展名。"""
    ext = _get_file_ext(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件类型: {ext}，仅支持 {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    return ext


def _extract_pdf_text(file_path: str) -> str:
    """从 PDF 提取文本（文本型 PDF）。"""
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            # 也提取表格文本
            tables = page.extract_tables() or []
            for table in tables:
                for row in table:
                    row_text = " | ".join(str(c) if c else "" for c in row)
                    if row_text.strip(" |"):
                        parts.append(row_text)
            if t.strip():
                parts.append(t)
    return "\n".join(parts)


def _extract_excel_text(file_path: str) -> str:
    """从 Excel 文件提取文本（所有 sheet）。"""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"=== Sheet: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(c) if c is not None else "" for c in row)
            if row_text.strip(" |"):
                parts.append(row_text)
    wb.close()
    return "\n".join(parts)


def _extract_docx_text(file_path: str) -> str:
    """从 Word 文档提取文本（段落 + 表格）。"""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    parts: list[str] = []

    # 段落
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n".join(parts)


def _extract_markdown_text(file_path: str) -> str:
    """Markdown 直接读取文本。"""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_text_from_file(file_path: str, filename: str) -> str:
    """根据文件类型提取文本。

    Args:
        file_path: 临时文件路径
        filename: 原始文件名（用于判断扩展名）

    Returns:
        提取的文本内容
    """
    ext = _validate_file(filename)

    if ext == ".pdf":
        text = _extract_pdf_text(file_path)
        if len(text.strip()) < 50:
            # 可能是扫描型 PDF，用 PyMuPDF 尝试提取
            import fitz

            doc = fitz.open(file_path)
            parts = []
            for page in doc:
                parts.append(page.get_text() or "")
            doc.close()
            text = "\n".join(parts)
        return text

    if ext in (".xlsx", ".xls"):
        return _extract_excel_text(file_path)

    if ext == ".docx":
        return _extract_docx_text(file_path)

    if ext in (".md", ".txt"):
        return _extract_markdown_text(file_path)

    raise ValueError(f"不支持的文件类型: {ext}")


def import_rules_from_document(
    db: Session,
    rule_set_id: uuid.UUID,
    file_content: bytes,
    filename: str,
) -> dict[str, Any]:
    """从上传的规则描述文档中提取文本，并调用 LLM 解析为结构化规则。

    Args:
        db: 数据库会话
        rule_set_id: 规则集 ID（导入规则归到该规则集下）
        file_content: 文件二进制内容
        filename: 原始文件名

    Returns:
        导入结果（与 import_rules_from_text 相同结构，额外含 extracted_text_preview）
    """
    ext = _validate_file(filename)

    # 写入临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(file_content)
        tmp_path = tmp.name

    try:
        # 提取文本
        logger.info("开始从文件 %s 提取规则文本", filename)
        text = extract_text_from_file(tmp_path, filename)

        if not text or not text.strip():
            raise ValueError("文件内容为空，无法提取规则文本")

        logger.info("文件 %s 提取到 %d 字符文本", filename, len(text))

        # 复用已有的文本导入逻辑
        result = import_rules_from_text(db, rule_set_id, text)
        result["extracted_text_preview"] = text[:500]  # 预览前 500 字符
        result["extracted_text_length"] = len(text)
        result["source_filename"] = filename
        return result

    finally:
        # 清理临时文件
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
